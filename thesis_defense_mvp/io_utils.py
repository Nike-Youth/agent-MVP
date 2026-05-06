from __future__ import annotations

import re
from pathlib import Path
from typing import Optional


def normalize_text(text: str) -> str:
    """清理文本中的多余空格和空行。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_txt(path: Path) -> str:
    """读取 txt / md 文件，兼容常见中文编码。"""
    encodings = ["utf-8", "utf-8-sig", "gbk", "ansi"]
    last_error: Optional[Exception] = None
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except Exception as exc:  # noqa: BLE001
            last_error = exc
    raise ValueError(f"无法读取文本文件编码：{path}。最后错误：{last_error}")


def read_docx(path: Path) -> str:
    """读取 docx 文件中的段落与表格文本。"""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("请先安装 python-docx：pip install python-docx") from exc

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"\n【表格 {table_index}】")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(normalize_text(cell.text))
            parts.append(" | ".join(row_text))

    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    """读取 PDF 文本。扫描件 PDF 不能稳定提取，需要先 OCR。"""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("请先安装 pypdf：pip install pypdf") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(f"\n\n【第 {page_index} 页】\n{text}")
    return "\n".join(parts)


def read_file(path_str: str | Path | None) -> str:
    """根据后缀读取 docx / pdf / txt / md 文件。"""
    if not path_str:
        return ""

    path = Path(path_str)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")

    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return normalize_text(read_txt(path))
    if suffix == ".docx":
        return normalize_text(read_docx(path))
    if suffix == ".pdf":
        return normalize_text(read_pdf(path))

    raise ValueError(f"暂不支持该文件类型：{suffix}。请使用 docx / pdf / txt / md。")


def save_uploaded_file(uploaded_file, target_dir: Path) -> Path:
    """Streamlit 上传文件保存工具。"""
    target_dir.mkdir(parents=True, exist_ok=True)
    file_path = target_dir / uploaded_file.name
    file_path.write_bytes(uploaded_file.getbuffer())
    return file_path
